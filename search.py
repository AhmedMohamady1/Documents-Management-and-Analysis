from pymongo import MongoClient
from gridfs import GridFS
from PyPDF2 import PdfReader
from docx import Document
import os
import datetime
import win32com.client
import urllib.parse
import fitz 
from pprint import pprint


client= MongoClient()
db = client.Project
collection = db["documents"]
fs = GridFS(db)

def insert_file(file):
    existing_file = collection.find_one({"name": file["name"]})
    if existing_file:
        print("File already exists in the database.")
    else:
        result = collection.insert_one(file)
        if result.inserted_id:
            print(f"File inserted successfully with ID: {result.inserted_id}")
        else:
            print("Failed to insert the File.")

# def count_pages_docx(docx_path):
#     # URL-decode the file path if needed
#     docx_path = urllib.parse.unquote(docx_path)
    
#     # Check if the file exists
#     if not os.path.exists(docx_path):
#         print(f"File not found: {docx_path}")
#         return

#     # Initialize Word application
#     word = win32com.client.Dispatch("Word.Application")
#     word.Visible = False  # Run Word in the background (set to False)

#     # Open the DOCX file
#     doc = word.Documents.Open(docx_path)
    
#     # Get the page count
#     page_count = doc.ComputeStatistics(2)  # 2 corresponds to WdStatisticPages (page count)

#     # Close the document and quit Word
#     doc.Close()
#     word.Quit()

#     return page_count

def count_pages(file_path):
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == ".pdf":
        with open(file_path, "rb") as pdf_file:
            pdf_reader = PdfReader(pdf_file)
            return len(pdf_reader.pages)
        
    elif file_extension == ".docx":
        doc = Document(file_path)
        page_count = sum(p.contains_page_break for p in doc.paragraphs) + 1
        return page_count

    elif file_extension == ".txt":
        return 1
    
    else:
        raise ValueError("Unsupported file format. Only .pdf and .docx files are supported.")

def extract_text_from_docx(file_path):
        text = ""
        document = Document(file_path)
        for paragraph in document.paragraphs:
            text += paragraph.text
        return text

def extract_text_from_pdf(file_path):
    with open(file_path, 'rb') as file:
        reader = PdfReader(file)
        text = ''
        for page in range(len(reader.pages)):
            text += reader.pages[page].extract_text()
        return text

def extract_text_from_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()


def extract_text_from_file(file_path):
    # Check if the file exists
    if not os.path.isfile(file_path):
        return f"File not found: {file_path}"
    
    # Get the file extension
    file_extension = os.path.splitext(file_path)[1].lower()
    
    try:
        # Handle PDF files
        if file_extension == ".pdf":
            return extract_text_from_pdf(file_path)
        
        # Handle DOCX files
        elif file_extension == ".docx":
            return extract_text_from_docx(file_path)
        
        # Handle TXT files
        elif file_extension == ".txt":
            return extract_text_from_txt(file_path)
        
        # Handle generic files
        else:
            return extract_text_from_generic(file_path)
    
    except Exception as e:
        return f"Error processing file: {e}"

def count_words_in_text(text):
    words = text.split()
    return len(words)

def count_characters_in_text(text, include_newlines=True):
    if not include_newlines:
        text = text.replace('\n', '')
    return len(text)

def Pull_File(file_path):
    text=extract_text_from_file(file_path)
    with open(file_path, "rb") as file:
        data = file.read()
        document = {
            "name": file_path.split("/")[-1],
            "contents": text,
            "file_data": data,
            "pages": count_pages(file_path),
            "words": count_words_in_text(text),
            "characters": count_characters_in_text(text),
            "modify date":datetime.datetime.fromtimestamp(os.path.getmtime(file_path)),
            "upload date":datetime.datetime.now()
        }
        insert_file(document)

def download_pdf(file_name, save_path):
    document = collection.find_one({"name": file_name})
    if not document:
        print(f"No file found with the name: {file_name}")
        return
    
    os.makedirs(os.path.dirname(save_path), exist_ok=True)
    
    with open(save_path, "wb") as file:
        file.write(document["file_data"])
    
    print(f"file downloaded successfully to {save_path}")


from nltk.corpus import stopwords
import nltk
from collections import Counter
import re
from datetime import datetime,timedelta

class Search:
    def __init__(self, mongodb_driver):
        self.__mongodb_driver=mongodb_driver

    def __db_query(self, search_terms, attributes='contents'):
        # Ensure attributes are valid
        valid_attributes = ['name', 'contents', 'modify date', 'upload date']
        
        # Convert single strings to lists
        if isinstance(search_terms, str):
            search_terms = [search_terms]
        if isinstance(attributes, str):
            attributes = [attributes]

        if not all(attr in valid_attributes for attr in attributes):
            return

        # Initialize the query
        search_queries = []
        for i in range(len(attributes)):
            if attributes[i] in ['upload date', 'modify date']:
                try:
                    # Parse date
                    date_parts = search_terms[i].split('-')
                    date = datetime(int(date_parts[0]), int(date_parts[1]), int(date_parts[2]))
                    start_of_day = date
                    end_of_day = date + timedelta(days=1)

                    # Add to search queries
                    search_queries.append({attributes[i]: {"$gte": start_of_day, "$lt": end_of_day}})
                except (ValueError, IndexError):
                    # Skip invalid date formats
                    continue
            else:
                # Add regex queries for text attributes
                search_queries.append({attributes[i]: {"$regex": search_terms[i], "$options": "i"}})

        # Combine queries with $or to match any condition
        combined_query = {"$and": search_queries}
        # Define the projection
        projection = {'_id': 0}

        # Perform the query
        return self.__mongodb_driver.find(combined_query, projection)

    
    def __print_document_metaData(self, document):
        print(f'Name: {document['name'].split('.')[0]}\nType: {document['name'].split('.')[-1]}\nNo. of Pages: {document['pages']}\nNo. of Words: {document['words']}\nNo. of Characters: {document['characters']}\nDate Uploaded: {str(document['upload date']).split('.')[0]}\nDate Modified: {str(document['modify date']).split('.')[0]}')
            
    def __most_common_words(self, words_list: list, n=5):
        word_counts = Counter(words_list)
        
        most_common = word_counts.most_common(n)
        
        print('\nMost commonly used words:\n')
        for i in range(len(most_common)):
            print(f'{i+1}) {most_common[i][0].title()} : {most_common[i][1]}')
            
    def __term_frequency(self, search_term, words):
        return round(words.count(search_term.lower())/len(words),6)
    
    def __content_preparation(self, content):
        words = self.__remove_unnecessary_words(content)
        for i in range(len(words)):
            words[i]=words[i].lower()
        return words
        
    def __remove_unnecessary_words(self,input_string):
        stop_words = set(stopwords.words('english'))
        words = re.findall(r'[A-Z][a-z]*|[^A-Z\s]+', input_string)
        filtered_words = [
        re.sub(r'[^a-zA-Z]', '', word)
        for word in words
        if word.lower() not in stop_words and re.sub(r'[^a-zA-Z]', '', word) and len(re.sub(r'[^a-zA-Z]', '', word)) > 1
    ]
        return filtered_words
        
    def extract_page_numbers_of_text(self, document: dict, search_term: str):
        pdf_data = document['file_data']
        pdf_file = fitz.open(stream=pdf_data, filetype="pdf")
        page_numbers = []

        for page_num in range(pdf_file.page_count):
            page = pdf_file.load_page(page_num)
            page_content = page.get_text("text")
            if search_term.lower() in page_content.lower():
                page_numbers.append(page_num + 1)
            
        return page_numbers

    def search_file(self, search_term: str, attribute: str):
        try:
            documents = list(self.__db_query(search_term, attribute))
            if len(documents)==0:
                print('Document not found.')
                return
            for document in documents:
                self.__print_document_metaData(document)
                self.__most_common_words(self.__remove_unnecessary_words(document['contents']))
                print('=================================================================')
        except Exception as e:
            print('Query error:',e)

    def search_contents(self, search_term):
        try:
            documents = list(self.__db_query(search_term))
            if len(documents)==0:
                print('Document not found.')
                return
            print(f'Documents that contain "{search_term.title()}"')
            for document in documents:
                self.__print_document_metaData(document)
                prepared_content=self.__content_preparation(document['contents'])
                print(f'Term Frequency of the term "{search_term}": {self.__term_frequency(search_term, prepared_content)} ({prepared_content.count(search_term.lower())}/{len(prepared_content)})')
                page_numbers = self.extract_page_numbers_of_text(document, search_term)
                print(f'Pages containing the search term "{search_term}": {", ".join(map(str, page_numbers))}')
                print('=================================================================')
        except Exception as e:
            print('Query error:',e)
    

if __name__=='__main__':
    # Pull_File("C:/Users/AhMeD/Desktop/C1_W1.pdf")
    search = Search(collection)
    search.search_contents("machine")