from pymongo import MongoClient
from gridfs import GridFS
from PyPDF2 import PdfReader
from docx import Document
import os
import datetime
import re
import win32com.client
import PyPDF2

class Mongo:
    def __init__(self):
        self.client= MongoClient()
        self.db= self.client.Project
        self.collection= self.db["documents"]
        self.fs = GridFS(self.db)

    def insert_file(self,file):
        existing_file = self.collection.find_one({"name": file["name"]})
        if existing_file:
            print("File already exists in the database.")
        else:
            result = self.collection.insert_one(file)
            if result.inserted_id:
                print(f"File inserted successfully with ID: {result.inserted_id}")
            else:
                print("Failed to insert the File.")

    def delete_file(self,file_name):
        filter_criteria = {"name": file_name}
        
        result = self.collection.delete_one(filter_criteria)
        if result.deleted_count > 0:
            print(f"File '{file_name}' deleted successfully.")
        else:
            print(f"No file matched the name '{file_name}'.")

    def get_word_page_count(self,word_data):
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False  
        
        temp_path = "temp_document.docx"
        with open(temp_path, "wb") as temp_file:
            temp_file.write(word_data)
        
        doc = word.Documents.Open(os.path.abspath(temp_path))

        page_count = doc.ComputeStatistics(2)
        
        doc.Close()
        word.Quit()
        os.remove(temp_path)
        
        return page_count


    def count_pages(self,file_path, file_data):
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == ".pdf":
            with open(file_path, "rb") as pdf_file:
                pdf_reader = PdfReader(pdf_file)
                return len(pdf_reader.pages)
            
        elif file_extension == ".docx":
            return self.get_word_page_count(file_data)

        elif file_extension == ".txt":
            return 1
        
        else:
            raise ValueError("Unsupported file format. Only .pdf and .docx files are supported.")

    def extract_text_from_docx(self,file_path):
            text = ""
            document = Document(file_path)
            for paragraph in document.paragraphs:
                text += paragraph.text
            return text

    def extract_text_from_pdf(self,file_path):
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            text = ''
            for page in range(len(reader.pages)):
                text += reader.pages[page].extract_text()
            return text

    def extract_text_from_txt(self,file_path):
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()

    def extract_text_from_generic(self,file_path):
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read()


    def extract_text_from_file(self,file_path):
        # Check if the file exists
        if not os.path.isfile(file_path):
            return f"File not found: {file_path}"
        
        # Get the file extension
        file_extension = os.path.splitext(file_path)[1].lower()
        
        try:
            # Handle PDF files
            if file_extension == ".pdf":
                return self.extract_text_from_pdf(file_path)
            
            # Handle DOCX files
            elif file_extension == ".docx":
                return self.extract_text_from_docx(file_path)
            
            # Handle TXT files
            elif file_extension == ".txt":
                return self.extract_text_from_txt(file_path)
            
            # Handle generic files
            else:
                return self.extract_text_from_generic(file_path)
        
        except Exception as e:
            return f"Error processing file: {e}"

    def get_word_word_count(self,word_data):
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False  
        
        temp_path = "temp_document.docx"
        with open(temp_path, "wb") as temp_file:
            temp_file.write(word_data)
        
        doc = word.Documents.Open(os.path.abspath(temp_path))

        word_count = doc.ComputeStatistics(0)  # Use 0 to count words
        
        doc.Close()
        word.Quit()
        os.remove(temp_path)
        
        return word_count
        
    def count_words_in_text(self,file_path,data):
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"The file {file_path} does not exist.")
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_extension == '.pdf':
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ''
                    for page in reader.pages:
                        text += page.extract_text()
                
                words = re.findall(r'\b\w+\b', text.lower())
                return len(words)
            
            elif file_extension == '.docx':
                return self.get_word_word_count(data)

            
            elif file_extension == '.txt':
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
                
                words = re.findall(r'\b\w+\b', text.lower())
                return len(words)
            
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
        
        except Exception as e:
            print(f"Error processing file {file_path}: {e}")
            return 0

    def get_character_count(self,word_data):
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False 
        
        temp_path = "temp_document.docx"
        
        with open(temp_path, "wb") as temp_file:
            temp_file.write(word_data)
        
        doc = word.Documents.Open(os.path.abspath(temp_path))

        char_count = doc.ComputeStatistics(5)  
        
        doc.Close()
        word.Quit()
        
        os.remove(temp_path)
        
        return char_count

    def count_characters(self,file_path,data):
        file_extension = file_path.split('.')[-1].lower()
        
        if file_extension == "pdf":
            return self.count_characters_in_pdf(file_path)
        elif file_extension == "docx":
            return self.get_character_count(data)
        elif file_extension == "txt":
            return self.count_characters_in_txt(file_path)
        else:
            raise ValueError("Unsupported file format. Only PDF, DOCX, and TXT are supported.")

    def count_characters_in_pdf(self,file_path):
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""  
        return len(text)

    def count_characters_in_docx(self,file_path):
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text
        return len(text)

    def count_characters_in_txt(self,file_path):
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        return len(text)

    def Pull_File(self,file_path):
        text=self.extract_text_from_file(file_path)
        with open(fr"{file_path}", "rb") as file:
            data = file.read()
            count_page=self.count_pages(file_path,data)
            count_word=self.count_words_in_text(file_path,data)
            count_character=self.count_characters(file_path,data)
            document = {
                "name": file_path.split("\\")[-1],
                "contents": text,
                "file_data": data,
                "pages": count_page,
                "words": count_word,
                "characters": count_character,
                "modify date":datetime.datetime.fromtimestamp(os.path.getmtime(file_path)),
                "upload date":datetime.datetime.now()
            }
            self.insert_file(document)

    def download_pdf(self,file_name, save_path):
        document = self.collection.find_one({"name": file_name})
        if not document:
            print(f"No file found with the name: {file_name}")
            return
        
        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        
        with open(save_path, "wb") as file:
            file.write(document["file_data"])
        
        print(f"file downloaded successfully to {save_path}")

if __name__=='__main__':
    mongo=Mongo()
    
